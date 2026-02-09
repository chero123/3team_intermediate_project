# 라이브러리
import os
import glob
import zlib
import struct
import olefile
import re
import json
import hashlib
import unicodedata
from datetime import datetime
from typing import Literal, Dict
import io
import subprocess
import tempfile
import torch
from PIL import Image, ImageFilter, ImageStat
from tqdm import tqdm
from transformers import AutoProcessor
import csv
import base64
import time
from dataclasses import dataclass, field
import shutil
from pathlib import Path

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    print("python-docx 설치 필요: pip install python-docx")
    DOCX_AVAILABLE = False
    DocxDocument = None

try:
    from openai import OpenAI

    OPENAI_AVAILABLE = True
except ImportError:
    print("OpenAI 설치 필요: pip install openai")
    OPENAI_AVAILABLE = False
    OpenAI = None

try:
    import fitz  # PyMuPDF
except ImportError:
    print("PyMuPDF 설치 필요: pip install pymupdf")
    fitz = None

try:
    import pdfplumber
except ImportError:
    print("pdfplumber 설치 필요: pip install pdfplumber")
    pdfplumber = None

# ============================================
# PDF 파서 설정 (기본값: pdfplumber)
# ============================================
# "pdfplumber": 텍스트 품질 좋음, 표 추출 강력 (느림)
# "fitz": 속도 빠름 (2-5배)
PDF_PARSER: Literal["pdfplumber", "fitz"] = "pdfplumber"


try:
    from qwen_vl_utils import process_vision_info
    from vllm import LLM, SamplingParams

    VLM_AVAILABLE = True
    VLM_IMPORT_ERROR = ""
except ImportError as e:
    VLM_AVAILABLE = False
    VLM_IMPORT_ERROR = str(e)
    process_vision_info = None
    LLM = None
    SamplingParams = None
    print(f"[VLM] 필수 라이브러리 임포트 실패: {VLM_IMPORT_ERROR}")
    print("[VLM] 'pip install vllm qwen_vl_utils' 설치가 필요합니다.")

# ============================================
# VLM 프로바이더 설정
# ============================================
# "none": VLM 미사용
# "qwen3": Qwen3-VL 로컬 모델 (이미지에서 텍스트 추출)
# "openai": OpenAI VLM (표만 마크다운으로 추출, 중복 방지)
VLM_PROVIDER: Literal["none", "qwen3", "openai"] = "none"

# ============================================
# OpenAI VLM 설정 (표 추출용)
# ============================================
OPENAI_VLM_MODEL: str = "gpt-4o-mini"
OPENAI_VLM_RETRY: int = 2
OPENAI_VLM_SLEEP_SEC: float = 1.0

OPENAI_TABLE_PROMPT: str = """
너는 문서 페이지에서 '표'만 추출하는 도우미다.

규칙:
- 보이는 표만 추출한다 (추측 금지)
- 표는 Markdown table로 변환한다
- 표가 없으면 {"tables": []} 만 반환한다
- 설명 문장은 쓰지 않는다
- JSON만 출력한다

출력 형식:
{
  "tables": [
    {
      "caption": "표 제목 (없으면 빈 문자열)",
      "markdown": "| ... |"
    }
  ]
}
"""

OPENAI_IMAGE_PROMPT: str = """
이미지에 있는 차트/그래프/도표의 핵심 정보를 추출해라.

규칙:
- 제목, 범례, 축 레이블을 포함한다
- 수치와 단위를 정확히 기록한다
- 기간/날짜 정보가 있으면 포함한다
- 간결하게 항목별로 정리한다
- 이미지에 정보가 없으면 빈 문자열만 반환한다
"""

_OPENAI_CLIENT_CACHE: Dict[str, object] = {}

# ============================================
# VLM (Qwen3-VL) 설정
# ============================================
QWEN3_VL_ENABLED: bool = True if VLM_AVAILABLE else False

# 중요: 다운로드한 모델 폴더의 절대 경로를 입력하세요.
# 예: "C:/Models/Qwen2-VL-7B-Instruct" 또는 "/home/user/models/qwen3-vl"
QWEN3_VL_MODEL_PATH: str = "./models/qwen2.5-vl-7b"

# Qwen3-VL (VLM) 설정
QWEN3_VL_PROMPT: str = (
    "이미지에 있는 표/그래프/지표의 핵심 수치와 단위를 항목별로 정리해라. "
    "표는 Markdown table 형식으로 변환해라. "
    "가능하면 제목, 기간, 범례 정보를 포함하고 숫자는 정확히 기록해라. "
    "출력은 Markdown 형식으로 작성해라."
)
QWEN3_VL_MAX_TOKENS: int = 512
QWEN3_VL_GPU_MEMORY_UTILIZATION: float = 0.8
QWEN3_VL_MAX_MODEL_LEN: int = 4096

# 이미지 필터링 임계값
QWEN3_VL_MIN_IMAGE_PIXELS: int = 10000  # 최소 픽셀 수 (100x100)
QWEN3_VL_MIN_VARIANCE: float = 100.0  # 최소 분산 (단색 필터)
QWEN3_VL_MIN_NONWHITE_RATIO: float = 0.1  # 최소 비흰색 비율
QWEN3_VL_MIN_EDGE_ENERGY: float = 0.01  # 최소 엣지 에너지
QWEN3_VL_DEDUPE_IMAGES: bool = True  # 중복 이미지 제거

_QWEN3_VL_CACHE: Dict[str, object] = {}

# 지원 바이너리 확장자
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".hwp", ".docx"}

# CSV 메타데이터 컬럼명
CSV_TEXT_FIELD = "텍스트"
CSV_FILENAME_FIELD = "파일명"

# ============================================
# 안전한 파일명 생성 함수
# ============================================


def safe_filename(
    original_name: str, suffix: str = "_parsed.txt", max_bytes: int = 180
) -> str:
    """
    OS 파일명 길이(바이트) 제한을 피하기 위해:
    - 원본 이름 정규화
    - 위험한 문자 제거
    - UTF-8 바이트 기준으로 자르기
    - 짧은 해시를 붙여 충돌 방지

    Args:
        original_name: 원본 파일명 (확장자 제외)
        suffix: 저장 파일 접미사 (기본값: "_parsed.txt")
        max_bytes: 최대 바이트 수 (기본값: 180)

    Returns:
        안전하게 변환된 파일명
    """
    name = unicodedata.normalize("NFC", original_name)

    # 확장자 제거 (혹시 .hwp/.pdf가 포함되어 있어도 안전하게)
    base = re.sub(r"\.(hwp|pdf)$", "", name, flags=re.IGNORECASE)

    # 파일명에 들어가면 곤란한 문자 제거/치환 (윈도우 금지문자 포함)
    base = re.sub(r"[\/\\\:\*\?\"\<\>\|\n\r\t]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip()

    # 충돌 방지용 짧은 해시
    h = hashlib.sha1(name.encode("utf-8")).hexdigest()[:10]

    # suffix 고려해서 base를 바이트 기준으로 자르기
    # 최종 파일명: "{base}__{hash}{suffix}"
    tail = f"__{h}{suffix}"
    budget = max_bytes - len(tail.encode("utf-8"))
    if budget < 20:
        budget = 20

    b = base.encode("utf-8")
    if len(b) > budget:
        base = b[:budget].decode("utf-8", errors="ignore").rstrip()

    return f"{base}__{h}{suffix}"


# ============================================
# 텍스트 정제 함수
# ============================================


def clean_text(text: str) -> str:
    """텍스트 정제: 불필요한 공백 및 과도한 줄바꿈 제거"""
    if not text:
        return ""
    # 깨진 문자 제거 (한글, 영문, 숫자, 공백, 로마숫자, 기본 문장부호 유지)
    text = re.sub(
        r'[^\uAC00-\uD7A3\u2160-\u217Fa-zA-Z0-9\s.,!?():\-\[\]<>~·%/@\'"_=+○●■□▶◆※|*#{}]',
        "",
        text,
    )
    # 연속 공백/줄바꿈 정리
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# ============================================
# DOCX 파싱 함수
# ============================================


def extract_text_from_docx(path: str) -> str:
    """
    docx 파일에서 텍스트를 추출합니다.

    Args:
        path: docx 파일 경로

    Returns:
        str: 추출된 텍스트
    """
    if not DOCX_AVAILABLE:
        print("python-docx가 설치되어 있지 않습니다.")
        return ""
    try:
        doc = DocxDocument(path)
        return clean_text("\n".join(p.text for p in doc.paragraphs if p.text))
    except Exception as e:
        print(f"'{path}' DOCX 파싱 중 오류 발생: {e}")
        return ""


# ============================================
# HWP 파싱 함수
# ============================================


def parse_hwp_all(file_path: str) -> str:
    """HWP 파일에서 전체 텍스트를 추출합니다. (BodyText 파싱)"""
    if not os.path.exists(file_path):
        print(f"'{file_path}' 파일을 찾을 수 없습니다.")
        return ""
    try:
        f = olefile.OleFileIO(file_path)

        # 압축 여부 확인
        header = f.openstream("FileHeader").read()
        is_compressed = header[36] & 1

        texts = []

        # 모든 BodyText 섹션 처리
        for entry in f.listdir():
            if entry[0] == "BodyText":
                data = f.openstream(entry).read()

                # 압축 해제
                if is_compressed:
                    try:
                        data = zlib.decompress(data, -15)
                    except:
                        pass

                # 레코드 파싱
                i = 0
                while i < len(data):
                    if i + 4 > len(data):
                        break

                    rec_header = struct.unpack_from("<I", data, i)[0]
                    rec_type = rec_header & 0x3FF
                    rec_len = (rec_header >> 20) & 0xFFF

                    # 확장 길이 처리
                    if rec_len == 0xFFF:
                        if i + 8 > len(data):
                            break
                        rec_len = struct.unpack_from("<I", data, i + 4)[0]
                        i += 8
                    else:
                        i += 4

                    if i + rec_len > len(data):
                        break

                    # HWPTAG_PARA_TEXT (67)
                    if rec_type == 67 and rec_len > 0:
                        text_data = data[i : i + rec_len]
                        try:
                            text = text_data.decode("utf-16le", errors="ignore")
                            # HWP 제어문자 제거 (0x00~0x1F 범위의 특수 제어코드)
                            cleaned = []
                            for char in text:
                                code = ord(char)
                                if code >= 32 or char in "\n\r\t":
                                    cleaned.append(char)
                                elif code in [13, 10]:  # CR, LF
                                    cleaned.append("\n")
                            text = "".join(cleaned).strip()
                            if text:
                                texts.append(text)
                        except:
                            pass

                    i += rec_len

        f.close()

        if texts:
            result = "\n".join(texts)
            result = clean_text(result)
            print(f"'{file_path}' 파일 파싱 성공!")
            return result
        else:
            print(f"'{file_path}' 텍스트 추출 실패")
            return ""

    except Exception as e:
        print(f"'{file_path}' 파일 파싱 중 오류 발생: {e}")
        return ""


# ============================================
# PDF 파싱 함수
# ============================================


def parse_pdf_with_pdfplumber(file_path: str) -> str:
    """pdfplumber를 이용한 PDF 텍스트 추출 (텍스트 품질 좋음, 페이지 마커 포함)"""
    if pdfplumber is None:
        print("pdfplumber가 설치되어 있지 않습니다.")
        return ""

    try:
        with pdfplumber.open(file_path) as pdf:
            print(f"  총 {len(pdf.pages)}페이지 추출 중... (pdfplumber)")
            parts = []
            for i, p in enumerate(pdf.pages):
                page_text = p.extract_text()
                if page_text:
                    parts.append(f"<!-- page: {i + 1} -->\n{page_text.strip()}")
                else:
                    parts.append(f"<!-- page: {i + 1} -->")
            if parts:
                result = clean_text("\n\n".join(parts))
                print(f"'{file_path}' 파일 파싱 성공!")
                return result
            else:
                print(f"'{file_path}' 텍스트 추출 실패")
                return ""
    except Exception as e:
        print(f"'{file_path}' PDF 파싱 중 오류 발생: {e}")
        return ""


def _is_inside_bbox(obj: dict, bbox: tuple, tolerance: float = 5.0) -> bool:
    """객체가 bbox 안에 있는지 확인 (tolerance 허용)"""
    if "x0" not in obj or "top" not in obj:
        return False
    x0, top, x1, bottom = bbox
    return (
        obj.get("x0", 0) >= x0 - tolerance
        and obj.get("top", 0) >= top - tolerance
        and obj.get("x1", 0) <= x1 + tolerance
        and obj.get("bottom", 0) <= bottom + tolerance
    )


def parse_pdf_exclude_tables(file_path: str) -> str:
    """
    pdfplumber로 표 영역을 제외하고 텍스트만 추출합니다.
    OpenAI VLM으로 표를 별도 추출할 때 중복을 방지합니다.
    """
    if pdfplumber is None:
        print("pdfplumber가 설치되어 있지 않습니다.")
        return ""

    try:
        with pdfplumber.open(file_path) as pdf:
            print(f"  총 {len(pdf.pages)}페이지 추출 중... (표 영역 제외)")
            parts = []

            for i, page in enumerate(pdf.pages):
                # 표 영역 감지
                tables = page.find_tables()
                table_bboxes = [table.bbox for table in tables]

                if table_bboxes:
                    # 표 영역을 제외한 텍스트 추출
                    filtered_page = page.filter(
                        lambda obj: not any(
                            _is_inside_bbox(obj, bbox) for bbox in table_bboxes
                        )
                    )
                    text = filtered_page.extract_text()
                else:
                    # 표가 없으면 전체 텍스트 추출
                    text = page.extract_text()

                if text:
                    parts.append(f"<!-- page: {i + 1} -->\n{text.strip()}")
                else:
                    parts.append(f"<!-- page: {i + 1} -->")

            if parts:
                result = clean_text("\n\n".join(parts))
                print(f"'{file_path}' 파일 파싱 성공! (표 영역 제외)")
                return result
            else:
                print(f"'{file_path}' 텍스트 추출 실패")
                return ""
    except Exception as e:
        print(f"'{file_path}' PDF 파싱 중 오류 발생: {e}")
        return ""


def parse_pdf_with_fitz(file_path: str) -> str:
    """PyMuPDF(fitz)를 이용한 PDF 텍스트 추출 (속도 빠름)"""
    if fitz is None:
        print("PyMuPDF가 설치되어 있지 않습니다.")
        return ""

    try:
        doc = fitz.open(file_path)
        results = []

        print(f"  총 {len(doc)}페이지 추출 중... (fitz)")

        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                results.append(f"<!-- page: {page_num + 1} -->\n{page_text.strip()}")
            else:
                results.append(f"<!-- page: {page_num + 1} -->")

        doc.close()

        if results:
            result = "\n\n".join(results)
            result = clean_text(result)
            print(f"'{file_path}' 파일 파싱 성공!")
            return result
        else:
            print(f"'{file_path}' 텍스트 추출 실패")
            return ""

    except Exception as e:
        print(f"'{file_path}' PDF 파싱 중 오류 발생: {e}")
        return ""


def _merge_vlm_into_pages(
    base_text: str, vlm_by_page: Dict[int, str], label: str = "tables"
) -> str:
    """
    페이지 마커가 포함된 본문 텍스트에 VLM 추출 결과를 인라인 삽입한다.

    <!-- page: N --> 마커를 찾아서 해당 페이지 텍스트 뒤에
    <!-- tables: start page N --> ... <!-- tables: end page N --> 블록을 삽입한다.

    Args:
        base_text: <!-- page: N --> 마커가 포함된 본문 텍스트
        vlm_by_page: {페이지번호: VLM 추출 텍스트}
        label: 삽입 블록 라벨 (기본값: "tables")

    Returns:
        str: VLM 결과가 인라인 삽입된 텍스트
    """
    if not vlm_by_page:
        return base_text

    lines = base_text.splitlines()

    # 페이지 마커 위치 수집
    page_starts = []
    for idx, line in enumerate(lines):
        m = re.match(r"^\s*<!--\s*page:\s*(\d+)\s*-->\s*$", line.strip())
        if m:
            page_starts.append((int(m.group(1)), idx))

    if not page_starts:
        return base_text

    # 페이지별 라인 범위 계산
    page_starts_sorted = sorted(page_starts, key=lambda x: x[1])
    page_to_end = {}
    for k in range(len(page_starts_sorted)):
        page_no, _ = page_starts_sorted[k]
        end_idx = (
            page_starts_sorted[k + 1][1]
            if k + 1 < len(page_starts_sorted)
            else len(lines)
        )
        page_to_end[page_no] = end_idx

    # 뒤에서부터 삽입 (인덱스 밀림 방지)
    for page_no in sorted(vlm_by_page.keys(), reverse=True):
        if page_no not in page_to_end:
            continue

        end_idx = page_to_end[page_no]
        vlm_content = vlm_by_page[page_no]

        insert_block = [
            "",
            f"<!-- {label}: start page {page_no} -->",
            vlm_content,
            f"<!-- {label}: end page {page_no} -->",
            "",
        ]

        lines[end_idx:end_idx] = insert_block

    return "\n".join(lines)


def record_vlm_log(file_path: str, vlm_dict: Dict[int, str], provider: str):
    """
    VLM 추출 결과를 JSON 로그 파일로 기록합니다.
    저장 위치: data/vlm_stats.json
    """
    log_dir = "data"
    os.makedirs(log_dir, exist_ok=True)
    log_path = os.path.join(log_dir, "vlm_stats.json")

    file_name = os.path.basename(file_path)
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "filename": file_name,
        "provider": provider,
        "pages_processed": list(vlm_dict.keys()),
        "total_images_extracted": sum(1 for text in vlm_dict.values() if text),
    }

    # 기존 로그 읽기
    logs = []
    if os.path.exists(log_path):
        try:
            with open(log_path, "r", encoding="utf-8") as f:
                logs = json.load(f)
        except:
            logs = []

    # 새 로그 추가 및 저장
    logs.append(entry)
    with open(log_path, "w", encoding="utf-8") as f:
        json.dump(logs, f, ensure_ascii=False, indent=2)
    print(f"[VLM LOG] '{file_name}' 추출 정보가 {log_path}에 기록되었습니다.")


def extract_text_from_pdf_with_vlm(path: str) -> tuple[str, str]:
    """
    PDF 텍스트와 VLM 이미지 텍스트를 함께 추출한다.
    VLM_PROVIDER 설정에 따라 OpenAI 또는 Qwen3-VL을 사용한다.
    VLM 결과는 해당 페이지 위치에 인라인 삽입된다.

    Args:
        path: PDF 경로

    Returns:
        tuple[str, str]: (VLM 인라인 삽입된 텍스트, "")
    """
    if VLM_PROVIDER == "openai":
        # OpenAI VLM: 표 영역 제외 텍스트 + 표/이미지 마크다운 추출
        base_text = parse_pdf_exclude_tables(path)
        table_dict = extract_tables_with_openai(path)
        image_dict = extract_images_with_openai(path)

        # 로그 기록 (표와 이미지 결과 합산)
        combined_dict = dict(table_dict)
        for p, t in image_dict.items():
            combined_dict[p] = combined_dict.get(p, "") + "\n\n" + t
        record_vlm_log(path, combined_dict, "openai")

        # 표를 페이지별 인라인 삽입
        merged = _merge_vlm_into_pages(base_text, table_dict, label="tables")
        # 이미지 분석도 페이지별 인라인 삽입
        merged = _merge_vlm_into_pages(merged, image_dict, label="images")
        return merged, ""

    elif VLM_PROVIDER == "qwen3" and QWEN3_VL_ENABLED:
        # Qwen3-VL: 기존 텍스트 + 이미지 VLM 추출 (페이지별 인라인 삽입)
        base_text = parse_pdf(path)
        vlm_dict = extract_vlm_text_from_pdf(path)

        # 로그 기록
        record_vlm_log(path, vlm_dict, "qwen3")

        merged = _merge_vlm_into_pages(base_text, vlm_dict, label="vlm")
        return merged, ""

    else:
        # VLM 미사용: 기존 텍스트 파싱만
        base_text = parse_pdf(path)
        return base_text, ""


def extract_text_from_pdf(path: str) -> str:
    """
    PDF 텍스트를 추출한다 (VLM_PROVIDER 설정에 따라 선택).

    Args:
        path: PDF 경로

    Returns:
        str: 추출된 텍스트 (마크다운 형식)
    """
    base_text, vlm_text = extract_text_from_pdf_with_vlm(path)
    parts = [t for t in [base_text, vlm_text] if t]
    if parts:
        return clean_text("\n\n".join(parts))
    return ""


def parse_pdf(file_path: str, parser: str = None) -> str:
    """
    PDF 파일에서 텍스트를 추출합니다.

    Args:
        file_path: PDF 파일 경로
        parser: "pdfplumber" 또는 "fitz" (None이면 전역 설정 사용)

    Returns:
        str: 추출된 텍스트
    """
    if not os.path.exists(file_path):
        print(f"'{file_path}' 파일을 찾을 수 없습니다.")
        return ""

    # 파서 선택 (매개변수 > 전역 설정)
    selected_parser = parser or PDF_PARSER

    if selected_parser == "pdfplumber":
        return parse_pdf_with_pdfplumber(file_path)
    else:
        return parse_pdf_with_fitz(file_path)


# ============================================
# OpenAI VLM 표 추출 함수
# ============================================


def _get_openai_client():
    """OpenAI 클라이언트를 지연 로딩한다."""
    if "client" in _OPENAI_CLIENT_CACHE:
        return _OPENAI_CLIENT_CACHE["client"]

    if not OPENAI_AVAILABLE:
        raise RuntimeError("OpenAI 패키지가 설치되어 있지 않습니다.")

    client = OpenAI()
    _OPENAI_CLIENT_CACHE["client"] = client
    return client


_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def _safe_json_loads(s: str) -> dict:
    """모델이 ```json ... ``` 같은 형식으로 줘도 최대한 JSON만 뽑아 파싱."""
    s = (s or "").strip()
    s = re.sub(r"^```json\s*|\s*```$", "", s, flags=re.IGNORECASE).strip()
    m = re.search(r"\{.*\}", s, flags=re.DOTALL)
    if m:
        s = m.group(0)
    # 제어문자 제거 (JSONDecodeError 방지)
    s = _CONTROL_CHARS_RE.sub("", s)
    return json.loads(s)


def _page_to_data_url(page, dpi: int = 140, img_format: str = "jpeg") -> str:
    """fitz 페이지를 base64 data URL로 변환 (JPEG 기본, 메모리 절약)"""
    pix = page.get_pixmap(dpi=dpi, alpha=False)
    if img_format.lower() == "png":
        raw = pix.tobytes("png")
        mime = "image/png"
    else:
        raw = pix.tobytes("jpeg")
        mime = "image/jpeg"
    del pix
    b64 = base64.b64encode(raw).decode("utf-8")
    return f"data:{mime};base64,{b64}"


def _extract_tables_from_page_openai(page, page_no: int, dpi: int = 140) -> list:
    """
    OpenAI Responses API + JSON 모드로 단일 페이지에서 표를 추출합니다.

    Returns:
        list: [{"caption": "...", "markdown": "| ... |"}, ...]
    """
    client = _get_openai_client()
    img_url = _page_to_data_url(page, dpi=dpi)

    last_err = None
    for attempt in range(OPENAI_VLM_RETRY + 1):
        try:
            # JSON 모드 시도 → 실패하면 일반 호출 fallback
            try:
                resp = client.responses.create(
                    model=OPENAI_VLM_MODEL,
                    response_format={"type": "json_object"},
                    input=[
                        {
                            "role": "system",
                            "content": [
                                {"type": "input_text", "text": OPENAI_TABLE_PROMPT}
                            ],
                        },
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "input_text",
                                    "text": f"{page_no}페이지에서 표만 추출해.",
                                },
                                {"type": "input_image", "image_url": img_url},
                            ],
                        },
                    ],
                )
                content = resp.output_text
            except Exception:
                # fallback: JSON 모드 없이 재시도
                resp = client.responses.create(
                    model=OPENAI_VLM_MODEL,
                    input=[
                        {
                            "role": "system",
                            "content": [
                                {"type": "input_text", "text": OPENAI_TABLE_PROMPT}
                            ],
                        },
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "input_text",
                                    "text": f"{page_no}페이지에서 표만 추출해.",
                                },
                                {"type": "input_image", "image_url": img_url},
                            ],
                        },
                    ],
                )
                content = resp.output_text

            result = _safe_json_loads(content)
            return result.get("tables", []) if isinstance(result, dict) else []
        except Exception as e:
            last_err = e
            time.sleep(OPENAI_VLM_SLEEP_SEC * (attempt + 1))

    print(f"[OpenAI VLM] 페이지 {page_no} 표 추출 실패: {last_err}")
    return []


def extract_tables_with_openai(file_path: str, dpi: int = 140) -> Dict[int, str]:
    """
    OpenAI VLM으로 PDF의 모든 페이지에서 표를 추출합니다.

    Args:
        file_path: PDF 파일 경로
        dpi: 이미지 렌더링 DPI

    Returns:
        Dict[int, str]: {페이지번호: 마크다운 표 텍스트}
    """
    if fitz is None:
        print("[OpenAI VLM] fitz 미설치 → 이미지 렌더링 불가")
        return {}

    if not OPENAI_AVAILABLE:
        print("[OpenAI VLM] OpenAI 패키지 미설치")
        return {}

    try:
        doc = fitz.open(file_path)
        print(f"[OpenAI VLM] PDF 표 추출 시작: {file_path} ({len(doc)}페이지)")

        result_dict: Dict[int, str] = {}
        for i in range(doc.page_count):
            page_no = i + 1
            page = doc.load_page(i)

            tables = _extract_tables_from_page_openai(page, page_no=page_no, dpi=dpi)

            if tables:
                page_parts = []
                for t in tables:
                    caption = (t.get("caption") or "").strip()
                    md_table = (t.get("markdown") or "").strip()

                    if caption:
                        page_parts.append(f"**[표] {caption}**")
                    if md_table:
                        page_parts.append(md_table)

                if page_parts:
                    result_dict[page_no] = "\n\n".join(page_parts)
                    print(f"  페이지 {page_no}: {len(tables)}개 표 추출")

        doc.close()

        if result_dict:
            print(f"[OpenAI VLM] 총 {len(result_dict)}개 페이지에서 표 추출 완료")
        else:
            print("[OpenAI VLM] 표가 발견되지 않았습니다.")
        return result_dict

    except Exception as e:
        print(f"[OpenAI VLM] 표 추출 오류: {e}")
        return {}


def _image_to_data_url(image: Image.Image, img_format: str = "jpeg") -> str:
    """PIL 이미지를 base64 data URL로 변환 (JPEG 기본, 메모리 절약)"""
    buffer = io.BytesIO()
    if img_format.lower() == "png":
        image.save(buffer, format="PNG")
        mime = "image/png"
    else:
        image.save(buffer, format="JPEG")
        mime = "image/jpeg"
    b64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return f"data:{mime};base64,{b64}"


def _analyze_image_with_openai(image: Image.Image, page_no: int, img_no: int) -> str:
    """
    OpenAI VLM으로 단일 이미지(차트/그래프)를 분석합니다.

    Args:
        image: PIL 이미지
        page_no: 페이지 번호
        img_no: 이미지 번호

    Returns:
        str: 추출된 텍스트 (빈 문자열 가능)
    """
    client = _get_openai_client()
    img_url = _image_to_data_url(image)

    last_err = None
    for attempt in range(OPENAI_VLM_RETRY + 1):
        try:
            resp = client.responses.create(
                model=OPENAI_VLM_MODEL,
                input=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": OPENAI_IMAGE_PROMPT},
                            {"type": "input_image", "image_url": img_url},
                        ],
                    },
                ],
            )
            content = resp.output_text
            return content.strip() if content else ""
        except Exception as e:
            last_err = e
            time.sleep(OPENAI_VLM_SLEEP_SEC * (attempt + 1))

    print(f"[OpenAI VLM] 페이지 {page_no} 이미지 {img_no} 분석 실패: {last_err}")
    return ""


def extract_images_with_openai(file_path: str) -> Dict[int, str]:
    """
    PDF 내 이미지를 추출하고, 의미 있는 이미지만 OpenAI VLM으로 분석합니다.
    Qwen3의 _is_meaningful_image() 필터링을 재사용합니다.

    Args:
        file_path: PDF 파일 경로

    Returns:
        Dict[int, str]: {페이지번호: 이미지 분석 텍스트}
    """
    if fitz is None:
        print("[OpenAI VLM] fitz 미설치 → 이미지 렌더링 불가")
        return {}

    if not OPENAI_AVAILABLE:
        print("[OpenAI VLM] OpenAI 패키지 미설치")
        return {}

    try:
        doc = fitz.open(file_path)
        print(f"[OpenAI VLM] PDF 이미지 분석 시작: {file_path} ({len(doc)}페이지)")

        result_dict: Dict[int, str] = {}
        seen_hashes: set[str] = set()

        for page_idx, page in enumerate(
            tqdm(doc, desc="[OpenAI VLM] 이미지 분석", unit="page")
        ):
            images = page.get_images(full=True)
            if not images:
                continue

            page_no = page_idx + 1
            page_parts = []

            for img_idx, img in enumerate(images):
                xref = img[0]
                base = doc.extract_image(xref)
                image_bytes = base.get("image")
                if not image_bytes:
                    continue

                # 중복 이미지 필터링 (Qwen3 방식 재사용)
                if QWEN3_VL_DEDUPE_IMAGES:
                    digest = hashlib.sha1(image_bytes).hexdigest()
                    if digest in seen_hashes:
                        continue
                    seen_hashes.add(digest)

                try:
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                except Exception:
                    continue

                # 의미 있는 이미지 필터링 (Qwen3 함수 재사용)
                if not _is_meaningful_image(image):
                    continue

                # OpenAI VLM으로 분석
                text = _analyze_image_with_openai(image, page_no, img_idx + 1)
                print(
                    f"[OpenAI VLM] page={page_no} image={img_idx + 1} "
                    f"text_len={len(text)}"
                )
                if text:
                    page_parts.append(f"#### [이미지 {img_idx + 1}]\n{text}")

            if page_parts:
                result_dict[page_no] = "\n\n".join(page_parts)

        doc.close()

        if result_dict:
            print(f"[OpenAI VLM] 총 {len(result_dict)}개 페이지에서 이미지 분석 완료")
        else:
            print("[OpenAI VLM] 분석할 이미지가 없습니다.")
        return result_dict

    except Exception as e:
        print(f"[OpenAI VLM] 이미지 분석 오류: {e}")
        return {}


# ============================================
# VLM (Qwen3-VL) 설정
# ============================================
# Qwen3-VL helpers
def _get_qwen3_vl():
    """
    Qwen3-VL 모델과 프로세서를 지연 로딩한다.

    Returns:
        tuple: (processor, model)
    """
    # 모델/프로세서는 무겁기 때문에 프로세스 전역 캐시로 재사용한다.
    if "processor" in _QWEN3_VL_CACHE and "llm" in _QWEN3_VL_CACHE:
        return _QWEN3_VL_CACHE["processor"], _QWEN3_VL_CACHE["llm"]

    # 로컬 체크포인트만 사용한다.
    model_source = QWEN3_VL_MODEL_PATH if os.path.exists(QWEN3_VL_MODEL_PATH) else None
    if model_source is None:
        raise RuntimeError(
            "Qwen3-VL 로컬 모델 경로를 찾을 수 없습니다. "
            f"필요 경로: {QWEN3_VL_MODEL_PATH}"
        )

    # Qwen3-VL은 processor에 이미지 전처리/토크나이즈 로직이 포함되어 있다.
    processor = AutoProcessor.from_pretrained(
        model_source,
        trust_remote_code=True,
        local_files_only=True,
    )
    # vLLM 엔진으로 FP8 모델을 로딩한다.
    llm = LLM(
        model=model_source,
        gpu_memory_utilization=QWEN3_VL_GPU_MEMORY_UTILIZATION,
        max_model_len=QWEN3_VL_MAX_MODEL_LEN,
        tensor_parallel_size=max(1, torch.cuda.device_count()),
        enforce_eager=False,
        seed=0,
        limit_mm_per_prompt={"image": 1, "video": 0},
    )
    _QWEN3_VL_CACHE["processor"] = processor
    _QWEN3_VL_CACHE["llm"] = llm
    return processor, llm


def _prepare_vllm_inputs(messages: list, processor: AutoProcessor) -> dict:
    """
    Qwen3-VL vLLM 입력 형식을 구성한다.

    Args:
        messages: Qwen3-VL 메시지 리스트
        processor: Qwen3-VL 프로세서

    Returns:
        dict: vLLM generate() 입력 dict
    """
    # Qwen3-VL은 chat template로 텍스트 프롬프트를 만든 뒤,
    # 이미지/비디오 입력을 별도로 멀티모달 데이터로 전달해야 한다.
    text = processor.apply_chat_template(
        messages, tokenize=False, add_generation_prompt=True
    )
    # 이미지/비디오 입력을 모델이 기대하는 텐서 형식으로 변환한다.
    image_inputs, video_inputs, video_kwargs = process_vision_info(
        messages,
        image_patch_size=processor.image_processor.patch_size,
        return_video_kwargs=True,
        return_video_metadata=True,
    )
    # vLLM의 multi_modal_data 포맷을 구성한다.
    mm_data = {}
    if image_inputs is not None:
        mm_data["image"] = image_inputs
    if video_inputs is not None:
        mm_data["video"] = video_inputs
    return {
        "prompt": text,
        "multi_modal_data": mm_data,
        "mm_processor_kwargs": video_kwargs,
    }


def _vlm_extract_from_image(image: Image.Image) -> str:
    """
    Qwen3-VL로 이미지에서 수치/표 정보를 텍스트로 추출한다.

    Args:
        image: PIL 이미지

    Returns:
        str: VLM 추출 텍스트 (빈 문자열 가능)
    """
    # Qwen3-VL 엔진/프로세서를 준비한다.
    processor, llm = _get_qwen3_vl()
    # Qwen3-VL은 이미지 입력을 메시지에 포함하고 vLLM 형식으로 변환한다.
    # Qwen3-VL은 "image + text"를 동일 메시지 안에 넣는다.
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image", "image": image},
                {"type": "text", "text": QWEN3_VL_PROMPT},
            ],
        }
    ]
    # vLLM이 요구하는 멀티모달 입력 dict로 변환한다.
    inputs = _prepare_vllm_inputs(messages, processor)
    try:
        # 수치 추출 목적이므로 온도 0으로 결정적 출력만 받는다.
        sampling_params = SamplingParams(
            temperature=0.0,
            max_tokens=QWEN3_VL_MAX_TOKENS,
            top_k=-1,
            stop_token_ids=[],
        )
        # vLLM은 입력 dict 리스트를 받는다.
        outputs = llm.generate([inputs], sampling_params=sampling_params)
        if not outputs:
            return ""
        text = outputs[0].outputs[0].text
        return text.strip() if text else ""
    except Exception as exc:
        print(f"[VLM] infer error: {exc}")
        return ""


def extract_vlm_text_from_pdf(path: str) -> Dict[int, str]:
    """
    PDF 페이지 이미지를 렌더링하고 Qwen3-VL로 텍스트를 추출한다.

    Args:
        path: PDF 경로

    Returns:
        Dict[int, str]: {페이지번호: VLM 추출 텍스트}
    """
    if fitz is None:
        print(f"[VLM] fitz 미설치 → 이미지 렌더링 불가: {path}")
        return {}
    if not QWEN3_VL_ENABLED:
        print(f"[VLM] disabled → 이미지 추출 스킵: {path}")
        return {}
    try:
        doc = fitz.open(path)
        print(f"[VLM] PDF images -> {path} (pages={len(doc)})")
        result_dict: Dict[int, str] = {}
        # 동일 이미지 중복 처리 방지를 위한 해시 캐시
        seen_hashes: set[str] = set()
        for page_index, page in enumerate(tqdm(doc, desc="[VLM] pages", unit="page")):
            images = page.get_images(full=True)
            if not images:
                continue
            page_no = page_index + 1
            page_parts = []
            for img_index, img in enumerate(images):
                xref = img[0]
                # PDF 내부 이미지 바이너리를 추출한다.
                base = doc.extract_image(xref)
                image_bytes = base.get("image")
                if not image_bytes:
                    continue
                # 중복 이미지(로고 반복 등)를 스킵한다.
                if QWEN3_VL_DEDUPE_IMAGES:
                    digest = hashlib.sha1(image_bytes).hexdigest()
                    if digest in seen_hashes:
                        continue
                    seen_hashes.add(digest)
                try:
                    # 바이너리를 PIL 이미지로 변환한다.
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                except Exception:
                    continue
                # 로고/단색/저정보 이미지 필터링
                if not _is_meaningful_image(image):
                    continue
                # VLM 추론 실행
                text = _vlm_extract_from_image(image)
                print(
                    f"[VLM] page={page_no} image={img_index + 1} "
                    f"text_len={len(text)}"
                )
                if text:
                    page_parts.append(f"#### [이미지 {img_index + 1}]\n{text}")
            if page_parts:
                result_dict[page_no] = "\n\n".join(page_parts)
        doc.close()
        print(f"[VLM] extracted images={len(result_dict)}")
        return result_dict
    except Exception as exc:
        print(f"[VLM] extract error: {exc}")
        return {}


def _is_meaningful_image(image: Image.Image) -> bool:
    """
    단순 로고/빈 이미지 등을 빠르게 걸러낸다.
    """
    # 너무 작은 이미지는 정보가 거의 없으므로 제외한다.
    width, height = image.size
    if width * height < QWEN3_VL_MIN_IMAGE_PIXELS:
        return False
    # 그레이스케일로 변환해 통계량을 계산한다.
    gray = image.convert("L")
    stats = ImageStat.Stat(gray)
    variance = stats.var[0] if stats.var else 0.0
    # 분산이 낮으면 단색/로고 가능성이 높다.
    if variance < QWEN3_VL_MIN_VARIANCE:
        return False
    # 거의 흰색(또는 거의 검은색)인 이미지인지 비율로 판별한다.
    hist = gray.histogram()
    total = max(1, sum(hist))
    nonwhite = sum(hist[:250]) / total
    if nonwhite < QWEN3_VL_MIN_NONWHITE_RATIO:
        return False
    # 엣지 에너지가 낮으면 텍스트/도표 정보가 부족하다고 판단한다.
    edges = gray.filter(ImageFilter.FIND_EDGES)
    edge_hist = edges.histogram()
    edge_energy = sum(i * count for i, count in enumerate(edge_hist)) / (total * 255.0)
    if edge_energy < QWEN3_VL_MIN_EDGE_ENERGY:
        return False
    return True


# ============================================
# HWP → PDF 변환 함수 (LibreOffice soffice)
# ============================================


def _convert_hwp_to_pdf(hwp_path: str, out_dir: str = None) -> str:
    """
    LibreOffice(soffice)를 이용하여 HWP 파일을 PDF로 변환한다.

    Args:
        hwp_path: HWP 파일 경로
        out_dir: PDF 출력 디렉토리 (None이면 임시 디렉토리 사용)

    Returns:
        str: 변환된 PDF 파일 경로 (실패 시 빈 문자열)
    """
    if out_dir is None:
        out_dir = tempfile.mkdtemp(prefix="hwp2pdf_")
    else:
        os.makedirs(out_dir, exist_ok=True)

    hwp_abs = os.path.abspath(hwp_path)
    out_abs = os.path.abspath(out_dir)

    # LibreOffice 프로필 충돌 방지용 임시 프로필
    profile_dir = tempfile.mkdtemp(prefix="lo_profile_")

    cmd = [
        "soffice",
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        f"-env:UserInstallation=file:///{profile_dir.replace(os.sep, '/')}",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        out_abs,
        hwp_abs,
    ]

    try:
        print(f"[HWP→PDF] 변환 시작: {os.path.basename(hwp_path)}")
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)

        # 변환된 PDF 경로 확인
        stem = os.path.splitext(os.path.basename(hwp_path))[0]
        pdf_path = os.path.join(out_abs, stem + ".pdf")

        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print(f"[HWP→PDF] 변환 성공: {pdf_path}")
            return pdf_path
        else:
            msg = (res.stderr or res.stdout or "").strip()
            print(f"[HWP→PDF] 변환 실패: {os.path.basename(hwp_path)} / {msg}")
            return ""
    except FileNotFoundError:
        print("[HWP→PDF] LibreOffice(soffice)가 설치되어 있지 않거나 PATH에 없습니다.")
        return ""
    except subprocess.TimeoutExpired:
        print(f"[HWP→PDF] 변환 타임아웃: {os.path.basename(hwp_path)}")
        return ""
    except Exception as e:
        print(f"[HWP→PDF] 변환 오류: {e}")
        return ""


# ============================================
# 확장자 자동 감지 함수
# ============================================


def load_file_content(file_path: str, use_vlm: bool = False) -> str:
    """
    확장자 자동 감지 및 추출

    Args:
        file_path: 파일 경로
        use_vlm: True이면 VLM_PROVIDER 설정에 따라 VLM 추출 수행
            - "none": VLM 미사용 (기존 텍스트 파싱만)
            - "qwen3": Qwen3-VL로 이미지에서 텍스트 추출
            - "openai": 표 영역 제외 텍스트 + OpenAI VLM으로 표/이미지 추출

    Returns:
        str: 추출된 텍스트
    """
    ext = file_path.split(".")[-1].lower()
    if ext == "hwp":
        # qwen3 또는 openai 모드일 경우 HWP -> PDF 변환 수행
        if use_vlm and (VLM_PROVIDER == "openai" or VLM_PROVIDER == "qwen3"):
            # HWP → PDF 변환 후 PDF 파이프라인 (페이지 마커 + VLM 추출)
            pdf_path = _convert_hwp_to_pdf(file_path)
            if pdf_path:
                print(f"[HWP] PDF 변환 성공, VLM 추출 시작 ({VLM_PROVIDER})")
                base_text, vlm_text = extract_text_from_pdf_with_vlm(pdf_path)
                parts = [t for t in [base_text, vlm_text] if t]
                # 임시 PDF 파일 삭제 (선택 사항)
                # try: os.remove(pdf_path)
                # except: pass
                return clean_text("\n\n".join(parts)) if parts else ""
            else:
                # PDF 변환 실패 시 기존 HWP 파싱 fallback
                print(f"[HWP→PDF] 변환 실패 → HWP 직접 파싱으로 fallback")
                return parse_hwp_all(file_path)
        else:
            return parse_hwp_all(file_path)
    elif ext == "pdf":
        if use_vlm and VLM_PROVIDER != "none":
            # VLM 사용: extract_text_from_pdf_with_vlm 통합 함수 사용
            base_text, vlm_text = extract_text_from_pdf_with_vlm(file_path)
            parts = [t for t in [base_text, vlm_text] if t]
            return clean_text("\n\n".join(parts)) if parts else ""
        else:
            # VLM 미사용: 기존 텍스트 파싱만
            return parse_pdf(file_path)
    elif ext == "docx":
        return extract_text_from_docx(file_path)

    return ""


# ============================================
# 전체 파일 처리 함수
# ============================================


def process_all_files(
    input_dir: str = "data/original_data",
    output_dir: str = "data/parsing_data",
    save_to_file: bool = True,
    use_vlm: bool = False,
) -> dict:
    """
    입력 디렉토리의 모든 HWP, PDF 파일을 파싱합니다.

    Args:
        input_dir: 원본 파일 디렉토리
        output_dir: 파싱 결과 저장 디렉토리
        save_to_file: True이면 파일로 저장, False이면 저장하지 않음
        use_vlm: True이면 PDF 이미지에서 VLM 텍스트도 추출

    Returns:
        dict: {파일명: 추출된텍스트} 형태의 딕셔너리
    """
    # 출력 폴더 생성 (저장할 경우에만)
    if save_to_file:
        os.makedirs(output_dir, exist_ok=True)

    # HWP, PDF, DOCX 파일 수집
    hwp_files = glob.glob(f"{input_dir}/*.hwp")
    pdf_files = glob.glob(f"{input_dir}/*.pdf")
    docx_files = glob.glob(f"{input_dir}/*.docx")
    all_files = hwp_files + pdf_files + docx_files

    print(
        f"발견된 파일: HWP {len(hwp_files)}개, PDF {len(pdf_files)}개, DOCX {len(docx_files)}개"
    )
    print(f"총 {len(all_files)}개 파일 변환 시작...")
    print()

    parsed_docs = {}
    mapping = []  # 원본파일명 ↔ 저장파일명 매핑 (RAG 메타데이터용)
    success_count = 0
    fail_count = 0

    for file_path in all_files:
        file_name = os.path.basename(file_path)
        doc_id = file_name.rsplit(".", 1)[0]
        print(f"--- {file_name} ---")

        # 확장자에 따라 파싱
        content = load_file_content(file_path, use_vlm=use_vlm)

        if content:
            print(f"[추출된 텍스트 미리보기]")
            print(content[:300])
            print("...\n")

            parsed_docs[file_name] = content

            # 파일 저장 (옵션) - safe_filename 적용
            # VLM 사용 시 .md, 아니면 .txt
            if save_to_file:
                use_md = use_vlm and VLM_PROVIDER != "none"
                ext = ".md" if use_md else ".txt"

                output_name = safe_filename(doc_id, suffix=f"_parsed{ext}")
                output_path = os.path.join(output_dir, output_name)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(content)
                print(f"저장 완료: {output_path} ({len(content)}자)\n")

                # 매핑 정보 기록
                mapping.append(
                    {
                        "original_filename": file_name,
                        "source_path": file_path,
                        "saved_filename": output_name,
                        "saved_path": output_path,
                        "chars": len(content),
                        "vlm_used": use_vlm and VLM_PROVIDER != "none",
                    }
                )

            success_count += 1
        else:
            print("텍스트 추출 실패\n")
            fail_count += 1

    print(f"\n{'='*50}")
    print(
        f"변환 완료: 성공 {success_count}개 / 실패 {fail_count}개 (총 {len(all_files)}개)"
    )
    if save_to_file:
        print(f"저장 위치: {output_dir}/")

        # 매핑 파일 저장 (원본 ↔ 저장 파일명 추적용)
        if mapping:
            mapping_path = os.path.join(output_dir, "parsed_mapping.json")
            with open(mapping_path, "w", encoding="utf-8") as f:
                json.dump(mapping, f, ensure_ascii=False, indent=2)
            print(f"매핑 저장: {mapping_path}")

    return parsed_docs


# ============================================
# Document 데이터 클래스 및 메타데이터 로드
# ============================================


@dataclass
class Document:
    """RAG 파이프라인용 문서 객체"""

    id: str = ""
    text: str = ""
    metadata: dict = field(default_factory=dict)


def load_metadata_csv(csv_path: str) -> Dict[str, dict]:
    """
    CSV 메타데이터를 파일명 기준으로 로드합니다.

    Args:
        csv_path: CSV 파일 경로

    Returns:
        Dict[str, dict]: {파일명: 메타데이터} 딕셔너리
    """
    if not os.path.exists(csv_path):
        return {}

    metadata_by_name: Dict[str, dict] = {}
    with open(csv_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            name = (
                row.get(CSV_FILENAME_FIELD)
                or row.get("filename")
                or row.get("file")
                or row.get("id")
            )
            if not name:
                continue
            metadata_by_name[name] = {k: v for k, v in row.items() if v is not None}
    return metadata_by_name


def normalize_metadata(row: dict) -> dict:
    """
    CSV 컬럼명을 표준 키로 정규화합니다.

    Args:
        row: 원본 메타데이터 행

    Returns:
        dict: 표준화된 메타데이터
    """
    return {
        "notice_id": row.get("공고 번호"),
        "notice_round": row.get("공고 차수"),
        "project_name": row.get("사업명"),
        "project_amount": row.get("사업 금액"),
        "issuer": row.get("발주 기관"),
        "publish_date": row.get("공개 일자"),
        "bid_start": row.get("입찰 참여 시작일"),
        "bid_end": row.get("입찰 참여 마감일"),
        "summary": row.get("사업 요약"),
        "file_type": row.get("파일형식"),
        "filename": row.get("파일명") or row.get("filename"),
    }


def load_documents(
    data_dir: str = "data/original_data",
    output_dir: str = "data/parsing_data",
    metadata_csv: str | None = None,
    use_vlm: bool = False,
    save_to_file: bool = True,
) -> list[Document]:
    """
    데이터 디렉토리에서 문서를 로드하고, 파싱 결과를 .md로 저장합니다.
    PDF/HWP/DOCX를 모두 지원하며, VLM 이미지 텍스트도 포함합니다.

    Args:
        data_dir: 원본 파일 디렉토리
        output_dir: 파싱 결과 저장 디렉토리
        metadata_csv: 메타데이터 CSV 경로 (None이면 미사용)
        use_vlm: True이면 VLM_PROVIDER에 따라 이미지 텍스트 추출
        save_to_file: True이면 .md 파일로 저장

    Returns:
        list[Document]: 파싱된 Document 리스트
    """
    print(f"[LOAD] data_dir={data_dir}")

    # 메타데이터 로드
    metadata_map = load_metadata_csv(metadata_csv) if metadata_csv else {}
    print(f"[LOAD] metadata_csv={metadata_csv} rows={len(metadata_map)}")

    # 출력 폴더 생성
    if save_to_file:
        os.makedirs(output_dir, exist_ok=True)

    documents: list[Document] = []
    mapping = []
    success_count = 0
    fail_count = 0

    # 디렉토리 순회
    for root, _, files in os.walk(data_dir):
        for filename in files:
            path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            # 지원되는 바이너리 확장자만 처리
            if ext not in SUPPORTED_BINARY_EXTENSIONS:
                continue

            print(f"[LOAD] file={path} ext={ext}")

            # CSV 메타데이터 가져오기
            metadata = dict(metadata_map.get(filename, {}))
            normalized = normalize_metadata(metadata)

            # 확장자별 파싱
            if ext == ".pdf":
                if use_vlm and VLM_PROVIDER != "none":
                    print(f"[PDF] start extract_text_with_vlm path={path}")
                    pdf_text, image_text = extract_text_from_pdf_with_vlm(path)
                    print(f"[VLM] image_text_len={len(image_text)} path={path}")
                    text = clean_text(
                        "\n\n".join([t for t in [pdf_text, image_text] if t])
                    )
                    print(f"[PDF] done text_len={len(text)} path={path}")
                    # VLM 이미지 텍스트를 메타데이터에 저장
                    if image_text:
                        metadata["vlm_image_text"] = image_text
                        metadata["vlm_image_text_present"] = True
                    else:
                        metadata["vlm_image_text_present"] = False
                else:
                    text = parse_pdf(path)
                    print(f"[PDF] text_len={len(text)} path={path}")
            elif ext == ".hwp":
                # qwen3 또는 openai 모드일 경우
                if use_vlm and (VLM_PROVIDER == "openai" or VLM_PROVIDER == "qwen3"):
                    pdf_path = _convert_hwp_to_pdf(path)
                    if pdf_path:
                        print(f"[HWP] PDF 변환 완료: {pdf_path}")
                        base_text, vlm_text = extract_text_from_pdf_with_vlm(pdf_path)
                        text = clean_text(
                            "\n\n".join([t for t in [base_text, vlm_text] if t])
                        )

                        # VLM 추출 여부 메타데이터 업데이트 (qwen3도 포함)
                        if vlm_text:
                            metadata["vlm_image_text"] = vlm_text
                            metadata["vlm_image_text_present"] = True
                        else:
                            metadata["vlm_image_text_present"] = False
                    else:
                        print(f"[HWP→PDF] 변환 실패 → HWP 직접 파싱으로 fallback")
                        text = parse_hwp_all(path)
                else:
                    text = parse_hwp_all(path)
                print(f"[HWP] text_len={len(text)} path={path}")
            else:
                text = ""

            if not text:
                print(f"텍스트 추출 실패: {filename}")
                fail_count += 1
                continue

            success_count += 1

            # 문서 ID 생성
            doc_id = os.path.relpath(path, data_dir)

            # 메타데이터 합치기
            metadata.update(normalized)
            metadata.update({"path": path, "filename": filename})

            # Document 객체 생성
            documents.append(Document(id=doc_id, text=text, metadata=metadata))

            # .md 파일 저장
            if save_to_file:
                doc_stem = filename.rsplit(".", 1)[0]
                output_name = safe_filename(doc_stem, suffix="_parsed.md")
                output_path = os.path.join(output_dir, output_name)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(text)
                print(f"저장 완료: {output_path} ({len(text)}자)")

                mapping.append(
                    {
                        "original_filename": filename,
                        "source_path": path,
                        "saved_filename": output_name,
                        "saved_path": output_path,
                        "chars": len(text),
                        "vlm_used": use_vlm and VLM_PROVIDER != "none",
                    }
                )

    print(f"\n{'='*50}")
    print(f"로드 완료: 성공 {success_count}개 / 실패 {fail_count}개")

    # 매핑 파일 저장
    if save_to_file and mapping:
        mapping_path = os.path.join(output_dir, "parsed_mapping.json")
        with open(mapping_path, "w", encoding="utf-8") as f:
            json.dump(mapping, f, ensure_ascii=False, indent=2)
        print(f"매핑 저장: {mapping_path}")

    return documents


# ============================================
# 메인 실행
# ============================================


def set_pdf_parser(parser: str):
    """PDF 파서를 설정합니다."""
    global PDF_PARSER
    PDF_PARSER = parser


def set_vlm_provider(provider: str):
    """
    VLM 프로바이더를 설정합니다.

    Args:
        provider: "none", "qwen3", "openai" 중 하나
    """
    global VLM_PROVIDER
    if provider not in ("none", "qwen3", "openai"):
        raise ValueError(f"지원하지 않는 VLM 프로바이더: {provider}")
    VLM_PROVIDER = provider


def set_vlm_config(enabled: bool = True, model_path: str = None):
    """
    Qwen3-VL 설정을 변경합니다. (하위 호환성 유지)

    Args:
        enabled: VLM 사용 여부
        model_path: 모델 경로 (None이면 기존 경로 유지)
    """
    global QWEN3_VL_ENABLED, QWEN3_VL_MODEL_PATH
    QWEN3_VL_ENABLED = enabled
    if model_path:
        QWEN3_VL_MODEL_PATH = model_path


def set_openai_vlm_config(model: str = None):
    """
    OpenAI VLM 설정을 변경합니다.

    Args:
        model: OpenAI 모델명 (None이면 기존 설정 유지)
    """
    global OPENAI_VLM_MODEL
    if model:
        OPENAI_VLM_MODEL = model


if __name__ == "__main__":
    start_time = datetime.now()

    # 경로 설정
    INPUT_DIR = "data/original_data"
    OUTPUT_DIR = "data/parsing_data"

    # PDF 파서 설정
    # "pdfplumber": 텍스트 품질 좋음, 표 추출 강력 (느림)
    # "fitz": 속도 빠름 (2-5배)
    set_pdf_parser("fitz")  # 기본값: pdfplumber, fiz가 이미지 렌더링에 좋다는 답변

    # ============================================
    # VLM 프로바이더 설정
    # ============================================
    # "none": VLM 미사용 (기존 텍스트 파싱만)
    # "qwen3": Qwen3-VL 로컬 모델 (이미지에서 텍스트 추출)
    # "openai": OpenAI VLM (표 영역 제외 텍스트 + 표만 마크다운 추출)
    set_vlm_provider("qwen3")  # 기본값: none

    # Qwen3-VL 설정 (provider="qwen3" 사용 시)
    # 모델 다운로드 방법:
    #   pip install hf_transfer
    #   export HF_HUB_ENABLE_HF_TRANSFER=1
    #   huggingface-cli download Qwen/Qwen2.5-VL-7B-Instruct --local-dir ./models/qwen2.5-vl-7b
    set_vlm_config(
        enabled=True,
        model_path=r"D:\01.project\ai06-level2-project\models\qwen2.5-vl-7b",
    )

    # OpenAI VLM 설정 (provider="openai" 사용 시)
    # 환경변수 OPENAI_API_KEY 설정 필요
    set_openai_vlm_config(model="gpt-5")

    # VLM 사용 여부
    USE_VLM = True  # True: VLM_PROVIDER에 따라 VLM 추출 수행

    print(f"PDF 파서: {PDF_PARSER}")
    print(f"VLM 프로바이더: {VLM_PROVIDER}")
    print(f"VLM 사용: {USE_VLM}")

    # 전체 파일 처리
    parsed_docs = process_all_files(
        input_dir=INPUT_DIR,
        output_dir=OUTPUT_DIR,
        save_to_file=True,
        use_vlm=USE_VLM,
    )

    # ============================================
    # 사용 예시
    # ============================================
    # 1. 파싱만 (저장 O): set_vlm_provider("none")
    # 2. Qwen3-VL 이미지 추출: set_vlm_provider("qwen3"), USE_VLM = True
    # 3. OpenAI 표 추출: set_vlm_provider("openai"), USE_VLM = True
    #    - 표 영역 제외한 텍스트 + VLM으로 표만 마크다운 추출 (중복 방지)
    # PDF 파서 변경: set_pdf_parser("fitz") (속도 우선)

    elapsed = datetime.now() - start_time
    minutes, seconds = divmod(int(elapsed.total_seconds()), 60)
    print(
        f"실행 문서: {os.path.basename(__file__)}, 완료 시간: {minutes:02d}분 {seconds:02d}초"
    )
