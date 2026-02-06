from __future__ import annotations

"""
문서 로딩/파싱/청킹 모듈

섹션 구성:
- 메타데이터 로딩/정규화
- 파서(HWP/PDF/DOCX)
- 문서 로딩
- 청킹
"""

import csv
import io
import os
import re
import hashlib
import struct
import unicodedata
import zlib
from typing import Dict, Iterable, List, Optional, Literal

import olefile

from langchain_text_splitters import RecursiveCharacterTextSplitter
from docx import Document as DocxDocument
from pypdf import PdfReader
from PIL import Image, ImageFilter, ImageStat
import torch
from tqdm import tqdm
from transformers import AutoProcessor
from vllm import LLM, SamplingParams
from qwen_vl_utils import process_vision_info

try:
    import fitz  # PyMuPDF
except ImportError:
    print("PyMuPDF 설치 필요: pip install pymupdf")
    fitz = None

try:
    import pdfplumber
except ImportError:
    print("pdfplumber 설치 필요: pip install pdfplumber")
    pdfplumber = None

from .config import RAGConfig
from .types import Chunk, Document, Metadata

# 바이너리 확장자
SUPPORTED_BINARY_EXTENSIONS = {".pdf", ".hwp", ".docx"}

# CSV 텍스트 컬럼명을 지정
CSV_TEXT_FIELD = "텍스트"
CSV_FILENAME_FIELD = "파일명"

# PDF 파서 선택 (pdfplumber: 품질 우선, fitz: 속도 우선)
PDF_PARSER: Literal["pdfplumber", "fitz"] = "pdfplumber"

# Qwen3-VL (VLM) 설정
QWEN3_VL_PROMPT = (
    "이미지에 있는 표/그래프/지표의 핵심 수치와 단위를 항목별로 정리해라. "
    "가능하면 제목, 기간, 범례 정보를 포함하고 숫자는 정확히 기록해라."
)
_QWEN3_VL_CACHE: Dict[str, object] = {}


def safe_filename(original_name: str, suffix: str = "_parsed.txt", max_bytes: int = 180) -> str:
    """
    OS 파일명 길이 제한을 피하기 위해 안전한 파일명을 만든다.

    Args:
        original_name: 원본 파일명 (확장자 제외)
        suffix: 저장 파일 접미사
        max_bytes: 최대 바이트 길이

    Returns:
        str: 안전하게 변환된 파일명
    """
    name = unicodedata.normalize("NFC", original_name)
    base = re.sub(r"\.(hwp|pdf|docx)$", "", name, flags=re.IGNORECASE)
    base = re.sub(r"[\/\\\:\*\?\"\<\>\|\n\r\t]+", " ", base)
    base = re.sub(r"\s+", " ", base).strip()
    h = hashlib.sha1(name.encode("utf-8")).hexdigest()[:10]
    tail = f"__{h}{suffix}"
    budget = max_bytes - len(tail.encode("utf-8"))
    if budget < 20:
        budget = 20
    b = base.encode("utf-8")
    if len(b) > budget:
        base = b[:budget].decode("utf-8", errors="ignore").rstrip()
    return f"{base}__{h}{suffix}"


def clean_text(text: str) -> str:
    """
    텍스트 정제: 불필요한 공백 및 깨진 문자를 정리한다.

    Args:
        text: 원본 텍스트

    Returns:
        str: 정제된 텍스트
    """
    if not text:
        return ""
    text = re.sub(
        r'[^\uAC00-\uD7A3\u2160-\u217Fa-zA-Z0-9\s.,!?():\-\[\]<>~·%/@\'"_=+○●■□▶◆※]',
        "",
        text,
    )
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


# Qwen3-VL helpers
def _get_qwen3_vl(config: RAGConfig):
    """
    Qwen3-VL 모델과 프로세서를 지연 로딩한다.

    Returns:
        tuple: (processor, model)
    """
    # 모델/프로세서는 무겁기 때문에 프로세스 전역 캐시로 재사용한다.
    if "processor" in _QWEN3_VL_CACHE and "llm" in _QWEN3_VL_CACHE:
        return _QWEN3_VL_CACHE["processor"], _QWEN3_VL_CACHE["llm"]

    # 로컬 체크포인트만 사용한다.
    model_source = config.qwen3_vl_model_path if os.path.exists(config.qwen3_vl_model_path) else None
    if model_source is None:
        raise RuntimeError(
            "Qwen3-VL 로컬 모델 경로를 찾을 수 없습니다. "
            f"필요 경로: {config.qwen3_vl_model_path}"
        )

    # Qwen3-VL은 processor에 이미지 전처리/토크나이즈 로직이 포함되어 있다.
    processor = AutoProcessor.from_pretrained(
        model_source,
        trust_remote_code=True,
        local_files_only=True,
    )
    # vLLM 엔진으로 FP8 모델을 로딩한다.
    llm = LLM(
        model=model_source,
        gpu_memory_utilization=config.qwen3_vl_gpu_memory_utilization,
        max_model_len=config.qwen3_vl_max_model_len,
        tensor_parallel_size=max(1, torch.cuda.device_count()),
        enforce_eager=False,
        seed=0,
        limit_mm_per_prompt={"image": 1, "video": 0},
    )
    _QWEN3_VL_CACHE["processor"] = processor
    _QWEN3_VL_CACHE["llm"] = llm
    return processor, llm


def _prepare_vllm_inputs(messages: list, processor: AutoProcessor) -> dict:
    """
    Qwen3-VL vLLM 입력 형식을 구성한다.

    Args:
        messages: Qwen3-VL 메시지 리스트
        processor: Qwen3-VL 프로세서

    Returns:
        dict: vLLM generate() 입력 dict
    """
    # Qwen3-VL은 chat template로 텍스트 프롬프트를 만든 뒤,
    # 이미지/비디오 입력을 별도로 멀티모달 데이터로 전달해야 한다.
    text = processor.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
    # 이미지/비디오 입력을 모델이 기대하는 텐서 형식으로 변환한다.
    image_inputs, video_inputs, video_kwargs = process_vision_info(
        messages,
        image_patch_size=processor.image_processor.patch_size,
        return_video_kwargs=True,
        return_video_metadata=True,
    )
    # vLLM의 multi_modal_data 포맷을 구성한다.
    mm_data = {}
    if image_inputs is not None:
        mm_data["image"] = image_inputs
    if video_inputs is not None:
        mm_data["video"] = video_inputs
    return {
        "prompt": text,
        "multi_modal_data": mm_data,
        "mm_processor_kwargs": video_kwargs,
    }


def _vlm_extract_from_image(image: Image.Image, config: RAGConfig) -> str:
    """
    Qwen3-VL로 이미지에서 수치/표 정보를 텍스트로 추출한다.

    Args:
        image: PIL 이미지

    Returns:
        str: VLM 추출 텍스트 (빈 문자열 가능)
    """
    # Qwen3-VL 엔진/프로세서를 준비한다.
    processor, llm = _get_qwen3_vl(config)
    # Qwen3-VL은 이미지 입력을 메시지에 포함하고 vLLM 형식으로 변환한다.
    # Qwen3-VL은 "image + text"를 동일 메시지 안에 넣는다.
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "image", "image": image},
                {"type": "text", "text": QWEN3_VL_PROMPT},
            ],
        }
    ]
    # vLLM이 요구하는 멀티모달 입력 dict로 변환한다.
    inputs = _prepare_vllm_inputs(messages, processor)
    try:
        # 수치 추출 목적이므로 온도 0으로 결정적 출력만 받는다.
        sampling_params = SamplingParams(
            temperature=0.0,
            max_tokens=config.qwen3_vl_max_tokens,
            top_k=-1,
            stop_token_ids=[],
        )
        # vLLM은 입력 dict 리스트를 받는다.
        outputs = llm.generate([inputs], sampling_params=sampling_params)
        if not outputs:
            return ""
        text = outputs[0].outputs[0].text
        return text.strip() if text else ""
    except Exception as exc:
        print(f"[VLM] infer error: {exc}")
        return ""


def _is_useful_vlm_text(text: str) -> bool:
    """
    VLM 출력이 의미 있는 수치/지표를 포함하는지 판단한다.

    - 숫자가 포함되지 않으면 무의미한 설명으로 간주한다.
    - "no data", "없음" 등 부정 표현만 있는 경우는 제외한다.
    """
    if not text:
        return False
    # 숫자가 하나도 없으면 표/수치 정보가 없다고 판단한다.
    if not re.search(r"\d", text):
        return False
    lowered = text.lower()
    # 명시적 "무데이터" 메시지는 제외한다.
    no_data_phrases = [
        "no data",
        "데이터 없음",
        "수치 없음",
        "표 없음",
        "그래프 없음",
        "지표 없음",
        "없음",
        "없습니다",
        "불가능",
    ]
    if any(phrase in lowered for phrase in no_data_phrases):
        return False
    return True


def _extract_pdf_images_with_vlm(path: str, config: RAGConfig) -> str:
    """
    PDF 페이지 이미지를 렌더링하고 Qwen3-VL로 텍스트를 추출한다.

    Args:
        path: PDF 경로

    Returns:
        str: 이미지 기반 추출 텍스트 (페이지별 결합)
    """
    if fitz is None:
        print(f"[VLM] fitz 미설치 -> 이미지 렌더링 불가: {path}")
        return ""
    if not config.qwen3_vl_enabled:
        print(f"[VLM] disabled -> 이미지 추출 스킵: {path}")
        return ""
    try:
        doc = fitz.open(path)
        print(f"[VLM] PDF images -> {path} (pages={len(doc)})")
        # 페이지 단위로 이미지 추출 결과를 누적한다.
        results = []
        # 동일 이미지 중복 처리 방지를 위한 해시 캐시
        seen_hashes: set[str] = set()
        for page_index, page in enumerate(tqdm(doc, desc="[VLM] pages", unit="page")):
            images = page.get_images(full=True)
            if not images:
                continue
            for img_index, img in enumerate(images):
                xref = img[0]
                # PDF 내부 이미지 바이너리를 추출한다.
                base = doc.extract_image(xref)
                image_bytes = base.get("image")
                if not image_bytes:
                    continue
                # 중복 이미지(로고 반복 등)를 스킵한다.
                if config.qwen3_vl_dedupe_images:
                    digest = hashlib.sha1(image_bytes).hexdigest()
                    if digest in seen_hashes:
                        continue
                    seen_hashes.add(digest)
                try:
                    # 바이너리를 PIL 이미지로 변환한다.
                    image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                except Exception:
                    continue
                # 로고/단색/저정보 이미지 필터링
                if not _is_meaningful_image(image, config):
                    continue
                # VLM 추론 실행
                text = _vlm_extract_from_image(image, config)
                # 수치/지표가 없는 출력은 버린다.
                if not _is_useful_vlm_text(text):
                    continue
                print(
                    f"[VLM] page={page_index + 1} image={img_index + 1} "
                    f"text_len={len(text)}"
                )
                if text:
                    results.append(
                        f"[PAGE {page_index + 1} IMAGE {img_index + 1}]\n{text}"
                    )
        doc.close()
        print(f"[VLM] extracted images={len(results)}")
        return clean_text("\n\n".join(results)) if results else ""
    except Exception as exc:
        print(f"[VLM] extract error: {exc}")
        return ""


def _is_meaningful_image(image: Image.Image, config: RAGConfig) -> bool:
    """
    단순 로고/빈 이미지 등을 빠르게 걸러낸다.
    """
    # 너무 작은 이미지는 정보가 거의 없으므로 제외한다.
    width, height = image.size
    if width * height < config.qwen3_vl_min_image_pixels:
        return False
    # 그레이스케일로 변환해 통계량을 계산한다.
    gray = image.convert("L")
    stats = ImageStat.Stat(gray)
    variance = stats.var[0] if stats.var else 0.0
    # 분산이 낮으면 단색/로고 가능성이 높다.
    if variance < config.qwen3_vl_min_variance:
        return False
    # 거의 흰색(또는 거의 검은색)인 이미지인지 비율로 판별한다.
    hist = gray.histogram()
    total = max(1, sum(hist))
    nonwhite = sum(hist[:250]) / total
    if nonwhite < config.qwen3_vl_min_nonwhite_ratio:
        return False
    # 엣지 에너지가 낮으면 텍스트/도표 정보가 부족하다고 판단한다.
    edges = gray.filter(ImageFilter.FIND_EDGES)
    edge_hist = edges.histogram()
    edge_energy = sum(i * count for i, count in enumerate(edge_hist)) / (total * 255.0)
    if edge_energy < config.qwen3_vl_min_edge_energy:
        return False
    return True


# Metadata loading/normalization
def load_metadata_csv(csv_path: str) -> Dict[str, Metadata]:
    """
    CSV 메타데이터를 파일명 기준으로 로드하는 함수

    Args:
        csv_path: CSV 파일 경로

    Returns:
        Dict[str, Metadata]: 파일명 기준 메타데이터 맵
    """
    if not os.path.exists(csv_path):
        # 파일이 없으면 빈 딕셔너리를 반환
        return {}

    # 결과 딕셔너리를 준비
    metadata_by_name: Dict[str, Metadata] = {}

    # CSV 파일을 열 때 나오는 이상한 특수 문자 방지를 위해 utf-8-sig 인코딩 사용
    with open(csv_path, "r", encoding="utf-8-sig") as f:
        # DictReader를 생성
        reader = csv.DictReader(f)
        for row in reader:
            name = row.get(CSV_FILENAME_FIELD) or row.get("filename") or row.get("file") or row.get("id")
            if not name:
                continue
            # None을 제거하고 저장
            metadata_by_name[name] = {k: v for k, v in row.items() if v is not None}
    return metadata_by_name


def normalize_metadata(row: Metadata) -> Metadata:
    """
    CSV 컬럼명을 표준 키로 정규화하는 함수

    Args:
        row: 원본 메타데이터 행

    Returns:
        Metadata: 표준화된 메타데이터
    """
    return {
        "notice_id": row.get("공고 번호"),
        "notice_round": row.get("공고 차수"),
        "project_name": row.get("사업명"),
        "project_amount": row.get("사업 금액"),
        "issuer": row.get("발주 기관"),
        "publish_date": row.get("공개 일자"),
        "bid_start": row.get("입찰 참여 시작일"),
        "bid_end": row.get("입찰 참여 마감일"),
        "summary": row.get("사업 요약"),
        "file_type": row.get("파일형식"),
        "filename": row.get("파일명") or row.get("filename"),
    }

# HWP parsing
def parse_hwp_all(path: str) -> str:
    """
    HWP 파일에서 BodyText를 파싱해 텍스트를 추출한다.

    Args:
        path: HWP 파일 경로

    Returns:
        str: 추출된 텍스트
    """
    if not os.path.exists(path):
        return ""
    try:
        # HWP는 OLE 컨테이너이므로 olefile로 스트림을 읽는다.
        f = olefile.OleFileIO(path)
        header = f.openstream("FileHeader").read()
        # FileHeader의 압축 플래그 확인
        is_compressed = header[36] & 1
        texts: List[str] = []
        for entry in f.listdir():
            if entry[0] != "BodyText":
                continue
            data = f.openstream(entry).read()
            if is_compressed:
                try:
                    # HWP BodyText는 zlib raw deflate 형식일 수 있다.
                    data = zlib.decompress(data, -15)
                except zlib.error:
                    pass
            i = 0
            while i < len(data):
                if i + 4 > len(data):
                    break
                # HWP 레코드 헤더 파싱 (type/length)
                rec_header = struct.unpack_from("<I", data, i)[0]
                rec_type = rec_header & 0x3FF
                rec_len = (rec_header >> 20) & 0xFFF
                if rec_len == 0xFFF:
                    if i + 8 > len(data):
                        break
                    rec_len = struct.unpack_from("<I", data, i + 4)[0]
                    i += 8
                else:
                    i += 4
                if i + rec_len > len(data):
                    break
                # HWPTAG_PARA_TEXT(67)만 추출 대상
                if rec_type == 67 and rec_len > 0:
                    text_data = data[i : i + rec_len]
                    try:
                        # HWP 텍스트는 UTF-16LE로 저장된다.
                        text = text_data.decode("utf-16le", errors="ignore")
                        cleaned_chars = []
                        for char in text:
                            code = ord(char)
                            if code >= 32 or char in "\n\r\t":
                                cleaned_chars.append(char)
                            elif code in [13, 10]:
                                cleaned_chars.append("\n")
                        text = "".join(cleaned_chars).strip()
                        if text:
                            texts.append(text)
                    except UnicodeDecodeError:
                        pass
                i += rec_len
        f.close()
        if texts:
            return clean_text("\n".join(texts))
        return ""
    except Exception:
        return ""


def extract_text_from_hwp(path: str) -> str:
    """
    HWP 텍스트를 추출한다.

    Args:
        path: HWP 경로

    Returns:
        str: 추출된 텍스트
    """
    return parse_hwp_all(path)


# PDF parsing
def parse_pdf_with_pdfplumber(path: str) -> str:
    """
    pdfplumber로 PDF 텍스트를 추출한다.

    Args:
        path: PDF 경로

    Returns:
        str: 추출된 텍스트
    """
    if pdfplumber is None:
        return ""
    try:
        # pdfplumber는 텍스트 품질이 좋은 대신 속도가 느릴 수 있다.
        with pdfplumber.open(path) as pdf:
            text = " ".join([p.extract_text() for p in pdf.pages if p.extract_text()])
        return clean_text(text) if text else ""
    except Exception:
        return ""


def parse_pdf_with_fitz(path: str) -> str:
    """
    PyMuPDF(fitz)로 PDF 텍스트를 추출한다.

    Args:
        path: PDF 경로

    Returns:
        str: 추출된 텍스트
    """
    if fitz is None:
        return ""
    try:
        # PyMuPDF는 속도가 빠르지만 텍스트 품질은 문서에 따라 달라진다.
        doc = fitz.open(path)
        results = []
        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                results.append(page_text.strip())
        doc.close()
        return clean_text("\n\n".join(results)) if results else ""
    except Exception:
        return ""

def extract_text_from_pdf(path: str, config: Optional[RAGConfig] = None) -> str:
    """
    PDF 텍스트를 추출한다 (parser 설정에 따라 선택).

    Args:
        path: PDF 경로

    Returns:
        str: 추출된 텍스트
    """
    config = config or RAGConfig()
    base_text, image_text = extract_text_from_pdf_with_vlm(path, config)
    if base_text or image_text:
        return clean_text("\n\n".join([t for t in [base_text, image_text] if t]))
    return ""


def extract_text_from_pdf_with_vlm(path: str, config: RAGConfig) -> tuple[str, str]:
    """
    PDF 텍스트와 VLM 이미지 텍스트를 함께 추출한다.

    Args:
        path: PDF 경로

    Returns:
        tuple[str, str]: (본문 텍스트, 이미지 텍스트)
    """
    # 설정된 파서를 먼저 사용한다.
    if PDF_PARSER == "pdfplumber":
        print(f"[PDF] parser=pdfplumber path={path}")
        text = parse_pdf_with_pdfplumber(path)
        base_text = text or ""
    elif PDF_PARSER == "fitz":
        print(f"[PDF] parser=fitz path={path}")
        text = parse_pdf_with_fitz(path)
        base_text = text or ""
    else:
        print(f"[PDF] parser=none path={path}")
        base_text = ""

    # 텍스트 파서 결과가 비어 있으면 pypdf로 fallback한다.
    try:
        if not base_text:
            print(f"[PDF] fallback=pypdf path={path}")
            reader = PdfReader(path)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            base_text = clean_text(text)
    except Exception:
        base_text = base_text or ""

    # PDF 내 이미지가 있을 수 있으므로 Qwen3-VL로 이미지 텍스트를 한 번만 추출한다.
    image_text = _extract_pdf_images_with_vlm(path, config)
    return base_text, image_text


# DOCX parsing
def extract_text_from_docx(path: str) -> str:
    """
    docx 파일에서 텍스트를 추출하는 함수

    Args:
        path: docs 파일 경로

    Returns:
        str: 추출된 텍스트
    """
    # docx 문서를 로드
    doc = DocxDocument(path)
    # 문단 텍스트를 결합
    return clean_text("\n".join(p.text for p in doc.paragraphs if p.text))


# Document loading
def load_documents(
    data_dir: str,
    metadata_csv: str | None = None,
    config: Optional[RAGConfig] = None,
) -> List[Document]:
    """
    데이터 디렉토리에서 문서를 로드하는 함수

    Args:
        data_dir: 데이터 디렉토리 경로
        metadata_csv: 메타데이터 CSV 경로

    Returns:
        List[Document]: 문서 리스트
    """
    # 설정 객체 준비
    config = config or RAGConfig()

    print(f"[LOAD] data_dir={data_dir}")
    # 메타데이터를 로드
    metadata_map = load_metadata_csv(metadata_csv) if metadata_csv else {}
    print(f"[LOAD] metadata_csv={metadata_csv} rows={len(metadata_map)}")

    # 결과 리스트를 준비
    documents: List[Document] = []
    # 디렉토리를 순회하며 지원 확장자만 파싱한다.
    for root, _, files in os.walk(data_dir):
        for filename in files:
            # 파일 경로를 구성
            path = os.path.join(root, filename)
            # 확장자를 추출
            ext = os.path.splitext(filename)[1].lower()
            if ext in SUPPORTED_BINARY_EXTENSIONS:
                print(f"[LOAD] file={path} ext={ext}")

            # 파일별 메타데이터를 가져온다
            # CSV 메타데이터를 복사해 문서 메타데이터에 합친다.
            metadata = dict(metadata_map.get(filename, {}))
            # 메타데이터를 정규화한다
            normalized = normalize_metadata(metadata)

            # 바이너리 문서면 SV 텍스트를 확인
            if ext in SUPPORTED_BINARY_EXTENSIONS:
                # PDF라면 PDF를 파싱
                if ext == ".pdf":
                    # PDF는 텍스트 + 이미지(VLM) 모두 추출한다.
                    print(f"[PDF] start extract_text_with_vlm path={path}")
                    pdf_text, image_text = extract_text_from_pdf_with_vlm(path, config)
                    print(f"[VLM] image_text_len={len(image_text)} path={path}")
                    text = clean_text("\n\n".join([t for t in [pdf_text, image_text] if t]))
                    print(f"[PDF] done extract_text_with_vlm text_len={len(text)} path={path}")
                    if image_text:
                        metadata["vlm_image_text"] = image_text
                        metadata["vlm_image_text_present"] = True
                    else:
                        metadata["vlm_image_text_present"] = False
                # HWP라면 HWP를 파싱
                elif ext == ".hwp":
                    text = extract_text_from_hwp(path)
                    print(f"[HWP] text_len={len(text)} path={path}")
                # docx라면 docx를 파싱
                elif ext == ".docx":
                    text = extract_text_from_docx(path)
                    print(f"[DOCX] text_len={len(text)} path={path}")
                else:
                    text = ""
            else:
                continue

            # 문서 ID를 만들기
            doc_id = os.path.relpath(path, data_dir)
            # 표준 메타데이터를 합친다
            metadata.update(normalized)
            # 경로 정보를 추가
            metadata.update({"path": path, "filename": filename})
            # 문서를 추가
            documents.append(Document(id=doc_id, text=text, metadata=metadata))
    return documents


# Chunking
def simple_chunk(text: str, chunk_size: int, overlap: int) -> List[str]:
    """
    RecursiveCharacterTextSplitter 기반 청킹을 수행

    Args:
        text: 입력 텍스트
        chunk_size: 청크 길이
        overlap: 청크 겹침 길이

    Returns:
        List[str]: 청크 문자열 리스트
    """
    # 입력 검증을 수행
    if chunk_size <= 0:
        return []

    # 문장/문단 경계를 우선하고, 부족하면 더 작은 단위로 재귀 분할한다.
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", "□", "。", ".", "!", "?", " ", ""],
    )
    return [c for c in splitter.split_text(text) if c.strip()]


def chunk_documents(docs: Iterable[Document], chunk_size: int, overlap: int) -> List[Chunk]:
    """
    문서 리스트를 청크 리스트로 변환하는 함수

    Args:
        docs: 문서
        chunk_size: 청크 길이
        overlap: 청크 겹침 길이

    Returns:
        List[Chunk]: 청크 리스트
    """
    all_chunks: List[Chunk] = []
    for doc in docs:
        for i, piece in enumerate(simple_chunk(doc.text, chunk_size, overlap)):
            # 청크 ID 생성
            chunk_id = f"{doc.id}::chunk_{i}"
            # 메타데이터를 복사
            metadata = dict(doc.metadata)
            # 청크 정보를 추가
            metadata.update({"chunk_index": i, "doc_id": doc.id})
            # 청크를 추가
            all_chunks.append(Chunk(id=chunk_id, text=piece, metadata=metadata))
    return all_chunks
